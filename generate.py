###################################################################################
#
# THESIS TITLE: EXTRACTING AND ORGANIZING DISASTER-RELATED PHILIPPINE COMMUNITY
#               RESPONSES FOR AIDING NATIONWIDE RISK REDUCTION PLANNING AND
#               RESPONSE (2020)
#
# AUTHOR/DEVELOPER: Nicco Nocon
# INSTITUTION: De La Salle University, Manila
# EMAIL: noconoccin@gmail.com or nicco_louis_nocon@dlsu.edu.ph
# CONTACT NUMBER: (+63) 917 819 9311
#
# SOURCE OF FUNDING: Philippine-California Advanced Research Institutes (PCARI)
#                    through Commission on Higher Education and Department of
#                    Science and Technology – Science Education Institute
#
###################################################################################
# REPORT GENERATION
# Transfers the extracted and organized information into a word file, for which
# acts as the report or final output of the software. Note that the text excel
# file can still be written for tracing or other form of output.
###################################################################################

import re
import webbrowser
import data_utils
from datetime import datetime
from docx import Document
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
# from docx.oxml.ns import qn
# from docx.shared import RGBColor

WNS_COLS_NUM = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num"


def add_timestamp(document):
    """
    Function for adding a timestamp. In Month-Date-Year Hours-Minutes-Seconds format (e.g., Oct-02-2019 18:51:50).

    Args:
        document (object): The instance of the document to be edited.
    """
    timestamp = document.add_paragraph()
    timestamp.paragraph_format.space_before = Pt(2)
    timestamp.add_run(str(datetime.now().strftime("%b-%d-%Y %H:%M:%S")), 'Emphasis')


def add_divider(document):
    """
    Function that adds a divider made from a 1x1 table object.

    Args:
        document (object): The instance of the document to be edited.
    """
    table = document.add_table(rows=1, cols=1)  # Table instance
    table.style = 'Light List Accent 3'  # Table Design/Formatting
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_text = table.cell(0, 0).paragraphs[0].add_run('_')  # Dummy cell text
    cell_text.font.size = Pt(3)  # Height Change
    cell_text.font.color.theme_color = MSO_THEME_COLOR_INDEX.ACCENT_3  # Color Change


def add_title(document, title_text):
    """
    Function that adds the title and formats it for display.

    Args:
         document (object): The instance of the document to be edited.
         title_text (str): The string input for the title label.
    """
    title = document.add_paragraph()  # Title paragraph instance
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Title Formatting
    title.paragraph_format.line_spacing = 1
    title.paragraph_format.space_after = Pt(2)
    title_text = title.add_run(title_text)  # Title text
    title_text.font.name = 'Segoe UI'  # Text Formatting
    title_text.font.size = Pt(22)
    title_text.font.bold = True


def set_document_margin(document, section_number, top, bottom, left, right):
    """
    Function for setting the document's margin (in inches). Used only on setting the first section.

    Args:
        document (object): The instance of the document to be edited.
        section_number (int): The section to be changed. -1 is for the first section.
        top (int): The number to be set on the top margin.
        bottom (int): The number to be set on the bottom margin.
        left (int): The number to be set on the left margin.
        right (int): The number to be set on the right margin.
    """
    document.sections[section_number].top_margin = Inches(top)
    document.sections[section_number].bottom_margin = Inches(bottom)
    document.sections[section_number].left_margin = Inches(left)
    document.sections[section_number].right_margin = Inches(right)


def set_number_of_page_columns(section, columns):
    """
    Function that sets the number of columns through xpath.

    Args:
        section (object): The section to be modified.
        columns (int): The number of columns to be applied.
    """
    section._sectPr.xpath("./w:cols")[0].set(WNS_COLS_NUM, str(columns))


def write_report(filename, malasakit_response_list, ranked_clusters_list):
    """
    Function that generates the report in word document (.docx) format. Default filename: Report.docx

    Args:
        filename (str): The file path of the word document.
        malasakit_response_list (list): The list containing the MalasakitResponse objects.
        ranked_clusters_list: The list containing the clustered and ranked insights.
    """
    document = Document()  # Create an instance of the document

    # Set Section Margin (Page Margin) into 1 inch all sides ('-1' is the current section)
    set_document_margin(document, section_number=-1, top=1, bottom=1, left=1, right=1)

    add_title(document, 'FILIPINO TEXT ANALYSIS TOOL REPORT')  # Title text
    add_divider(document)  # Divider
    add_timestamp(document)  # Timestamp
    # paragraph.add_run('dolor').bold = True

    document.add_paragraph("The information below were extracted and organized automatically.")  # Intro

    # Make a new section for a double column format
    section = document.add_section(WD_SECTION_START.CONTINUOUS)
    set_number_of_page_columns(section, 2)

    footer = section.footer.add_paragraph()
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer.paragraph_format.line_spacing = 1
    footer_text = footer.add_run('Extracting and Organizing Disaster-related Philippine Community Responses for Aiding '
                                 'Nationwide Risk Reduction Planning and Response (N. Nocon, 2019)')
    footer_text.font.name = 'Segoe UI'
    footer_text.font.size = Pt(10)

    cluster_ctr = 0
    for cluster in ranked_clusters_list:
        if isinstance(cluster, str):  # 'If element is a str == true' means that index 0 is the category string not id
            # Category
            if cluster_ctr != 0:
                document.add_paragraph().paragraph_format.space_after = Pt(0)  # Added for padding
            add_divider(document)
            category = document.add_paragraph()
            category.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Title Formatting
            category.paragraph_format.line_spacing = 1
            category.paragraph_format.space_before = Pt(8)
            category.paragraph_format.space_after = Pt(8)
            category_text = category.add_run(cluster.upper())  # Category text
            category_text.font.size = Pt(12)  # Text Formatting
            category_text.font.bold = True
            cluster_ctr = 0  # Reset index
            add_divider(document)
        else:
            # Body
            cluster_ctr += 1  # Entry/Cluster 1, ...
            entry = document.add_paragraph()
            entry.paragraph_format.line_spacing = 1  # Text formatting
            entry.paragraph_format.space_before = Pt(5)
            entry.paragraph_format.space_after = Pt(5)
            entry.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            entry_text = entry.add_run('Entry ' + str(cluster_ctr))
            entry_text.font.bold = True

            label_toggle = 0
            for element in cluster:  # Print contents per row. Might be better to bold labels
                if label_toggle == 0:
                    label = document.add_paragraph()
                    label_text = label.add_run('ID/s: ')
                    label_text.font.bold = True
                elif label_toggle == 1:
                    label = document.add_paragraph()
                    label_text = label.add_run('Frequency: ')
                    label_text.font.bold = True
                elif label_toggle == 2:
                    label = document.add_paragraph()
                    label_text = label.add_run('Proposed action: ')
                    label_text.font.bold = True
                elif label_toggle == 3:
                    label = document.add_paragraph()
                    label_text = label.add_run('Target: ')
                    label_text.font.bold = True
                # elif label_toggle == 4:
                #     label = document.add_paragraph()
                #     label_text = label.add_run('Location: ')
                #     label_text.font.bold = True
                label.paragraph_format.line_spacing = 1  # Text formatting
                label.paragraph_format.space_after = Pt(2)
                label.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                label.add_run(str(element))  # Content
                # label_text.font.bold = False
                label_toggle += 1  # Toggle for correct element label

    # Make a new section for a single column format
    section = document.add_section(WD_SECTION_START.NEW_PAGE)
    set_number_of_page_columns(section, 1)

    # A part of the report shows Malasakit Responses
    add_title(document, 'MALASAKIT RESPONSES REFERENCE LIST')
    add_divider(document)
    document.add_paragraph()  # Padding

    # Make a new section for a double column format
    section = document.add_section(WD_SECTION_START.CONTINUOUS)
    set_number_of_page_columns(section, 2)

    table = document.add_table(rows=len(malasakit_response_list)+1, cols=2)  # Table instance
    table.style = 'Light List Accent 3'  # Table Design/Formatting
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, 0).text = 'ID'
    table.cell(0, 1).text = 'Response'

    row_ctr = 1
    for response in malasakit_response_list:
        table.cell(row_ctr, 0).text = str(response.response_id)  # Response ID
        table.cell(row_ctr, 1).text = response.response  # Response
        row_ctr += 1

    for cell in table.columns[0].cells:
        cell._tc.tcPr.tcW.type = 'auto'

    document.save(filename)


def generate_web(web_filename, excel_filename, organization_format, cluster_limit):
    """
    Function that generates the report in a web/HTML format. It takes the contents of an excel file and displays the
    indicated number of entries for organize all entries and organize per response categories. Target words are linked
    to their respective Malasakit responses.

    Args:
        web_filename (str): The file path of the generated HTML file.
        excel_filename (str): The file path of the excel file to be converted into web version.
        organization_format (str): The organizational format used (i.e., OAE or ORC).
        cluster_limit (int): The limited number of clusters to be shown (Top 5, 10, 25). 0 value for unlimited.
    """
    print('GENERATING WEB VERSION')
    malasakit_response_list = []
    limit_counter = 0
    y = 0

    malasakit_response_list = data_utils.read_excel(excel_filename)  # Reference for Malasakit Responses

    clusters_list, category_count = data_utils.read_clusters_excel(excel_filename, organization_format)
    # print(clusters_list)

    # Generate website
    code = ""
    html_skeleton = open('web/html_skeleton.txt', 'r')  # Get HTML template
    string_list = html_skeleton.readlines()  # Put in the list the string contents per new line

    # Composing code to write on HTML file
    # Go through each line in the template
    if category_count == 0:
        category_count += 1

    for line in string_list:
        code += line
        # Place list of clusters after Cluster List header
        if 'Cluster List</h2>' in line:
            code += "<hr/>"

            for x in range(category_count):
                while y < len(clusters_list):
                    cluster = clusters_list[y]
                    # print(cluster)
                    # print(y)
                    if cluster_limit != 0:
                        limit_counter += 1
                        if organization_format == 'ORC':
                            if limit_counter == cluster_limit * 2 + 2:  # Check if it exceeded the limit for display
                                while not str(clusters_list[y]).isupper() and not y + 1 == len(clusters_list):
                                    if y + 1 != len(clusters_list):
                                        y += 1
                                break
                        elif limit_counter == cluster_limit * 2 + 1:  # Check if it exceeded the limit for display
                            break

                    if organization_format == 'ORC' and str(cluster).isupper():
                        code += "<h4 class=\"alert alert-success text-center\">" + cluster + "</h4>\n"
                    else:
                        if "Cluster" in cluster:
                            code += "<hr/><h5>" + cluster + "</h5><hr/>\n"
                        else:  # Tuple
                            code += "<p>Frequency: " + str(cluster[1]) + "<br>" + "Proposed Action: " + cluster[2] + \
                                    "<br>" + "Targets:"

                            # Hyperlink Assignment and Generation
                            # Compare target words with the response ids in the cluster,
                            # if the response contains the word, assign hyperlink to that cluster
                            response_id_list = cluster[0].split('|')  # Get each id in sentence ids

                            # Get lexicalized words
                            lexicalized_word_list = re.findall("\w+ \(\w+[, \w+]*\)", cluster[3])

                            tuple_to_string = ''.join(cluster[3])

                            # Remove lexicalized words
                            for word in lexicalized_word_list:
                                tuple_to_string = tuple_to_string.replace(word + ',', '')
                                tuple_to_string = tuple_to_string.replace(word, '')

                            tuple_to_string = re.sub(" +", " ", tuple_to_string)
                            tuple_to_string = tuple_to_string.strip()
                            tuple_to_string = tuple_to_string.strip(',')

                            target_list = tuple_to_string.split(', ')  # Get each word in target

                            # print(response_id_list)
                            # print(target_list)

                            # Writing and making lexicalized words hyperlink
                            for lex in lexicalized_word_list:  # For every sub-clusters
                                lex = re.sub("[(,)]", "", lex)  # Remove punctuations
                                word_list = lex.split(' ')
                                lex_count = 1
                                for word in word_list:  # For every word
                                    for response_id in response_id_list:  # For every response
                                        if lex_count == 1:
                                            if word in malasakit_response_list[int(response_id) - 1].response:
                                                code += " <a class=\"text-success\" href=\"#R" + response_id + "\">" + word + "</a> ("
                                                break
                                        elif lex_count > 1 and lex_count != len(word_list):
                                            if word in malasakit_response_list[int(response_id) - 1].response:
                                                code += "<a class=\"text-success\" href=\"#R" + response_id + "\">" + word + "</a>, "
                                                break
                                        else:
                                            if word in malasakit_response_list[int(response_id) - 1].response:
                                                code += "<a class=\"text-success\" href=\"#R" + response_id + "\">" + word + "</a>),"
                                                break
                                    lex_count += 1

                            # Writing and making target hyperlink
                            for target in target_list:
                                # print("TARGET: ", target)
                                for response_id in response_id_list:
                                    # print("Response ID: ", int(response_id)-1)
                                    # print("COMPARE: ", target in malasakit_response_list[int(response_id)-1].response)
                                    if target in malasakit_response_list[int(response_id)-1].response:
                                        code += " <a class=\"text-success\" href=\"#R" + response_id + "\">" + target + "</a>,"
                                        break
                                    # No case if not found - all should have link since they came from the data
                            code = code.strip(',')
                            code += "</p>"

                            code += "<ul class=\"list-group\">" + \
                                    "<li class=\"list-group-item list-group-item-success py-2\"><b>Sample Cluster Members</b></li>"

                            sample_counter = 0
                            for response_id in response_id_list:
                                if sample_counter == 5:
                                    break
                                else:
                                    code += "<li class=\"list-group-item py-2\"><b>Response " + response_id + ":</b> " + \
                                            malasakit_response_list[int(response_id)-1].response + "</li>"
                                sample_counter += 1
                            code += "</ul><br>"

                    y += 1
                limit_counter = 0

        # Place list of responses after Malasakit Response List header
        if 'Malasakit Response List</h2>' in line:
            code += "<hr/>"
            code += "<ul class=\"list-group list-group-flush\">"
            for response in malasakit_response_list:
                code += "<li class=\"list-group-item py-2\">" + "<p class=\"mb-0\" id=\'R" + str(response.response_id) + "\'><b>Response #" + \
                        str(response.response_id) + ":</b> " + response.response + "</p></li>"
            code += "</ul>"

    html_skeleton.close()

    f = open(web_filename, 'w+')

    message = code

    f.write(message)
    f.close()

    # Change path to reflect file location
    webbrowser.open_new_tab(web_filename)

# 'test/MalasakitResponses_small.xlsx'
# generate_web('web/index.html', 'test/Experiments/MalasakitResponses_934_dice_org_category.xlsx', 'ORC', 10)
